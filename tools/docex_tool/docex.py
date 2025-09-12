# tools/docex_tool/docex.py
# AANGEPAST: commentaar per belangrijke component / functie toegevoegd (Nederlands)

import os
import io
import tempfile
import re
import json
import streamlit as st
from groq import Groq
import docx
from docx.enum.text import WD_BREAK

from . import steps  # keep top-level import for helper functions that use steps
from core.docx_utils import read_docx, apply_replacements


# -----------------------------
# Helper: Groq client ophalen
# -----------------------------

def get_groq_client():
    """
    Maak en retourneer een Groq-client gebaseerd op een API key.

    Werking:
    - Probeert eerst om de API key uit de omgeving (GROQ_API_KEY) te halen.
    - Als die niet aanwezig is, zoekt het naar `.streamlit/secrets.toml` of
      in de huidige werkmap en leest `st.secrets` (Streamlit secrets) uit.
    - Als er geen sleutel gevonden wordt, toont het een fout in de Streamlit
      sidebar en stopt de app (`st.stop()`).

    Retour:
    - Een Groq(...) client object als alles goed gaat.
    - Roept st.stop() aan bij fouten of ontbrekende config (zodat de app niet
      door blijft lopen zonder geldige client).
    """
    api_key = os.environ.get("GROQ_API_KEY", "").strip()
    if api_key:
        try:
            return Groq(api_key=api_key)
        except Exception as e:
            st.sidebar.error(f"❌ Fout bij verbinden met Groq API (env): {e}")
            st.stop()
    possible = [
        os.path.expanduser("~/.streamlit/secrets.toml"),
        os.path.join(os.getcwd(), ".streamlit", "secrets.toml"),
    ]
    api_key = ""
    if any(os.path.exists(p) for p in possible):
        try:
            api_key = st.secrets.get("groq", {}).get("api_key", "").strip()
        except Exception:
            api_key = ""
    if not api_key:
        # Gebruiksvriendelijke foutmelding in de sidebar
        st.sidebar.error(
            "❌ Groq API key niet gevonden. Zet GROQ_API_KEY als env var of maak `.streamlit/secrets.toml` met [groq] api_key = \"...\""
        )
        st.stop()
    try:
        return Groq(api_key=api_key)
    except Exception as e:
        st.sidebar.error(f"❌ Fout bij verbinden met Groq API: {e}")
        st.stop()


# -----------------------------
# Parser: model-output -> python data
# -----------------------------

def parse_groq_json_array(content: str):
    """
    Probeer een JSON-array te parsen uit een (soms rommelige) string die het
    model teruggeeft.

    Acties:
    - Verwijdert numerieke keys als "1: { ... }" die soms uit modellen komen.
    - Probeert het eerste '[' ... ']' te isoleren en die JSON te laden.
    - Als JSON decode faalt, probeert het heuristieken: zoekt naar paired
      "find" en "replace" lines en bouwt daaruit een lijst van vervangingen.

    Retour:
    - Een Python-lijst met dicts (bijv. [{"find": "X", "replace": "Y"}, ...]).
    """
    cleaned = re.sub(r"\d+\s*:\s*{", "{", content)
    start, end = cleaned.find("["), cleaned.rfind("]") + 1
    json_str = cleaned[start:end] if start != -1 and end != -1 else cleaned
    try:
        return json.loads(json_str)
    except json.JSONDecodeError:
        # Fallback heuristiek: zoek "find" en "replace" paren in tekst
        repls = []
        lines = cleaned.splitlines()
        for i, ln in enumerate(lines):
            if '"find"' in ln:
                fm = re.search(r'"find"\s*:\s*"([^"]*)"', ln)
                rm = None
                if fm:
                    for nxt in lines[i + 1 :]:
                        m = re.search(r'"replace"\s*:\s*"([^"]*)"', nxt)
                        if m:
                            rm = m.group(1)
                            break
                if fm and rm:
                    repls.append({"find": fm.group(1), "replace": rm})
        return repls


# -----------------------------
# Core: vraag model om vervangingen
# -----------------------------

def get_replacements_from_model(groq_client: Groq, template_text: str, context_text: str):
    """
    Stuur TEMPLATE en CONTEXT naar het model en verwacht een JSON-array met
    objecten {find, replace} terug.

    Details:
    - Bouwt een eenvoudige prompt met TEMPLATE en CONTEXT.
    - Roept `groq_client.chat.completions.create` aan met een low temperature.
    - Verwacht dat het model alleen de JSON-array teruggeeft (system-role
      instrueert dat).
    - Parst de content met `parse_groq_json_array`.
    - Filtert eindresultaat op geldige 'find' items en verwijdert identieke
      find==replace regels.

    Retour:
    - Lijst met vervangingen (kan leeg zijn).
    """
    prompt = (
        "Gegeven TEMPLATE en CONTEXT, lever JSON-array met objecten {find, replace}."
        f"\n\nTEMPLATE:\n{template_text}\n\nCONTEXT:\n{context_text}"
    )
    try:
        resp = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            temperature=0.2,
            messages=[
                {"role": "system", "content": "Antwoord alleen JSON-array, geen extra tekst."},
                {"role": "user", "content": prompt},
            ],
        )
    except Exception as e:
        st.error(f"Fout bij model-aanroep: {e}")
        return []
    content = resp.choices[0].message.content
    repls = parse_groq_json_array(content)
    # Filter onbruikbare of noop vervangingen
    return [r for r in repls if r.get("find") and r.get("find") != r.get("replace")]


# -----------------------------
# Document bewerkingen
# -----------------------------

def apply_replacements_to_doc_and_bytes(doc_path: str, replacements: list[dict], include_changes_overview: bool = True) -> bytes:
    """
    Laad een .docx bestand, pas tekstvervangingen toe en retourneer de
    gewijzigde inhoud als bytes (klaar voor download).

    Werking:
    - Loopt over alle paragraph-run combinaties en tabellen in het document.
    - Voor elk stuk tekst wordt alle gevonden "find" strings vervangen door
      hun "replace" waarde.
    - Om styled runs te behouden: zet de samengevoegde tekst in de eerste run
      en leegt de overige runs.
    - (optioneel) Voegt aan het einde een overzicht toe met welke vervangingen
      zijn uitgevoerd en welke stappen werden gedaan (via `steps.get_steps()`).

    Retour:
    - Bytes van het aangepaste document (in-memory buffer).
    """
    doc = docx.Document(doc_path)

    def repl(runs):
        # Helper die een lijst van runs (deelstukken van een paragraaf) als
        # één tekst behandelt, vervangt en weer inruns stopt.
        if not runs:
            return
        txt = "".join(r.text for r in runs)
        for rp in replacements:
            txt = txt.replace(rp["find"], rp["replace"])
        runs[0].text = txt
        for r in runs[1:]:
            r.text = ""

    # Pas vervangingen toe in paragrafen
    for p in doc.paragraphs:
        repl(p.runs)
    # Pas vervangingen toe in tabelcellen
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    repl(p.runs)

    # Optioneel: voeg een overzicht van wijzigingen toe als laatste pagina
    if include_changes_overview:
        try:
            p_br = doc.add_paragraph()
            run = p_br.add_run()
            run.add_break(WD_BREAK.PAGE)
        except Exception:
            # Niet-cruciale stap — sommige templates of docx-versies kunnen
            # dit niet toelaten; we falen stil
            pass
        try:
            doc.add_paragraph("Aangepaste onderdelen", style="Heading 1")
        except Exception:
            doc.add_paragraph("Aangepaste onderdelen")
        for rp in replacements:
            para = doc.add_paragraph(f"• {rp['find']} → {rp['replace']}")
            try:
                para.style = "List Bullet"
            except Exception:
                pass
        step_lines = steps.get_steps() if hasattr(steps, "get_steps") else []
        if step_lines:
            try:
                doc.add_paragraph("Uitgevoerde stappen", style="Heading 2")
            except Exception:
                doc.add_paragraph("Uitgevoerde stappen")
            for s in step_lines:
                doc.add_paragraph(f"• {s}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# -----------------------------
# Veilig text uitlezer (simpler): alleen plain text uit docx halen
# -----------------------------

def _safe_read_docx_text(path: str) -> str:
    """
    Probeer op een veilige manier de tekst uit een .docx te lezen.
    Retourneert lege string bij fouten.
    """
    try:
        doc = docx.Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        return ""


# -----------------------------
# Streamlit UI - hoofdfunctie
# -----------------------------

def run(show_nav: bool = True):
    """
    Entrypoint voor de Streamlit-app.

    Verantwoordelijkheden:
    - Initialiseren van stappen-tracker (import defensief om circular imports te vermijden).
    - Configureren van de pagina-UI (layout, knoppen, uploaders).
    - Inputs: template (.docx) en context (.docx of .txt).
    - Vraagt het model om vervangingen aan (`get_replacements_from_model`).
    - Toont gevonden vervangingen, genereert het aangepaste document en
      biedt een download-button aan.

    De functie is rijk aan Streamlit UI-logica en is bedoeld om interactief te draaien.
    """
    # defensive import to avoid NameError / circular import issues
    try:
        from . import steps as _steps_mod
    except Exception:
        import importlib

        _steps_mod = importlib.import_module("tools.docex_tool.steps")

    # use the local reference for internal calls
    steps_mod = _steps_mod
    steps_mod.clear_steps()

    # allow outer app to configure page; safe to keep
    st.set_page_config(page_title="DOCX Generator", layout="wide", initial_sidebar_state="expanded")
    st.markdown(
        """
        <style>
        .stButton>button, .stDownloadButton>button {font-size:18px; font-weight:bold; padding:0.6em 1.2em;}
        .big-header {font-size:2.5rem; font-weight:bold; margin-bottom:0.3em;}
        .section-header {font-size:1.75rem; font-weight:600; margin-top:1em; margin-bottom:0.5em;}
        .stTextArea textarea {font-family:monospace;}
        </style>
        """,
        unsafe_allow_html=True,
    )

    groq_client = get_groq_client()
    steps_mod.record_step("Groq client aangemaakt")

    # show internal nav only when requested; otherwise default to Generator
    if show_nav:
        page = st.sidebar.radio("🔖 Navigatie", ("Home", "Generator", "Info"), index=0)
    else:
        page = "Generator"

    if page == "Home":
        # Introductie / uitleg tab
        st.markdown("<div class='big-header'>🏠 Welkom bij de DOCX Generator</div>", unsafe_allow_html=True)
        st.markdown(
            """
            Gebruik deze tool om snel **Word-templates** bij te werken met **nieuwe context**.
            
            - Ga naar **Generator**
            - Upload je **template** en **context**
            - Klik op **Genereer aangepast document**
            - Download en behoud je opmaak!
            """,
            unsafe_allow_html=True,
        )

    elif page == "Generator":
        # De belangrijkste gebruikersinterface: uploaden, preview en genereren
        st.markdown("<div class='big-header'>🚀 Generator</div>", unsafe_allow_html=True)
        col1, col2 = st.columns(2)
        tpl_path = None
        context = ""
        with col1:
            st.markdown("<div class='section-header'>📄 Template Upload</div>", unsafe_allow_html=True)
            tpl_file = st.file_uploader("Kies .docx template", type="docx", key="tpl")
            if tpl_file:
                # Sla het ge-uploade template tijdelijk op zodat we de inhoud kunnen lezen
                tpl_path_dir = tempfile.mkdtemp()
                tpl_path = os.path.join(tpl_path_dir, "template.docx")
                with open(tpl_path, "wb") as f:
                    f.write(tpl_file.getbuffer())
                tpl_text = _safe_read_docx_text(tpl_path)
                st.text_area("Template-inhoud", tpl_text, height=250, key="tpl_pre")
                steps_mod.record_step("Template geüpload")
        with col2:
            st.markdown("<div class='section-header'>📝 Context Upload</div>", unsafe_allow_html=True)
            ctx_file = st.file_uploader("Kies .docx/.txt context", type=["docx", "txt"], key="ctx")
            if ctx_file:
                tmp_c = tempfile.mkdtemp()
                if ctx_file.type and ctx_file.type.endswith("document"):
                    # Context in .docx formaat — bewaar en lees de tekst
                    cpath = os.path.join(tmp_c, "context.docx")
                    with open(cpath, "wb") as f:
                        f.write(ctx_file.getbuffer())
                    context = _safe_read_docx_text(cpath)
                    steps_mod.record_step("Context (.docx) geüpload")
                else:
                    # Plain text context
                    context = ctx_file.read().decode("utf-8", errors="ignore")
                    steps_mod.record_step("Context (.txt) geüpload")
                st.text_area("Context-inhoud", context, height=250, key="ctx_pre")

        # Sidebar toont de eerdere stappen voor traceerbaarheid
        st.sidebar.markdown("### Uitgevoerde stappen")
        for s in steps_mod.get_steps():
            st.sidebar.markdown(f"- {s}")

        if tpl_path and context:
            st.markdown("---")
            if st.button("🎉 Genereer aangepast document"):
                # Gebruiker start de generatieflow
                steps_mod.record_step("Start generatie")
                tpl_text = _safe_read_docx_text(tpl_path)
                replacements = get_replacements_from_model(groq_client, tpl_text, context)
                steps_mod.record_step(f"Model returned {len(replacements)} replacement(s)")
                st.markdown("<div class='section-header'>✨ Aangepaste onderdelen</div>", unsafe_allow_html=True)
                if replacements:
                    for rp in replacements:
                        st.write(f"• **{rp['find']}** → **{rp['replace']}**")
                else:
                    st.info("Geen vervangingen gevonden (leeg resultaat of model gaf niks bruikbaars terug).")
                out_bytes = apply_replacements_to_doc_and_bytes(tpl_path, replacements)
                steps_mod.record_step("Document gegenereerd (bytes klaar)")
                st.markdown("<div class='section-header'>📝 Volledige stappen</div>", unsafe_allow_html=True)
                for s in steps_mod.get_steps():
                    st.write(f"- {s}")
                st.download_button(
                    "⬇️ Download aangepast document",
                    data=out_bytes,
                    file_name="aangepast_template.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
        else:
            st.info("Upload eerst template en context om te starten.")
    else:
        # Info tab: gebruikersadvies en tips
        st.markdown("<div class='big-header'>ℹ️ Info & Tips</div>", unsafe_allow_html=True)
        st.markdown(
            """
            **Tips voor optimaal gebruik:**
            - Zorg voor unieke, duidelijke tekstfragmenten.
            - Houd context-bestanden kort en concreet.
            - Controleer altijd de uiteindelijke output.
            - Voor complexe documenten kun je secties apart bijwerken.
            """,
            unsafe_allow_html=True,
        )


if __name__ == "__main__":
    run()
