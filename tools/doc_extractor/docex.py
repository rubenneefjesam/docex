# tools/docex_tool/docex.py

import os
import io
import tempfile
import re
import json
import streamlit as st
from groq import Groq
import docx
from docx.enum.text import WD_BREAK

# Importeer interne modules
from . import steps  # steps.py zit in dezelfde map als deze file
from core.utils.docx_utils import read_docx, apply_replacements  # core/docx_utils.py

# -----------------------------
# Groq-client ophalen
# -----------------------------
def get_groq_client():
    api_key = os.environ.get("GROQ_API_KEY", "").strip()
    if api_key:
        try:
            return Groq(api_key=api_key)
        except Exception as e:
            st.sidebar.error(f"‚ùå Fout bij verbinden met Groq API (env): {e}")
            st.stop()
    possible = [
        os.path.expanduser("~/.streamlit/secrets.toml"),
        os.path.join(os.getcwd(), ".streamlit", "secrets.toml"),
    ]
    api_key = ""
    if any(os.path.exists(p) for p in possible):
        try:
            api_key = st.secrets.get("groq", {}).get("api_key", "").strip()
        except Exception:
            api_key = ""
    if not api_key:
        st.sidebar.error(
            "‚ùå Groq API key niet gevonden. Zet GROQ_API_KEY als env var of maak `.streamlit/secrets.toml` met [groq] api_key = \"...\""
        )
        st.stop()
    try:
        return Groq(api_key=api_key)
    except Exception as e:
        st.sidebar.error(f"‚ùå Fout bij verbinden met Groq API: {e}")
        st.stop()

# -----------------------------
# Parser: model-output -> vervangingen
# -----------------------------
def parse_groq_json_array(content: str):
    cleaned = re.sub(r"\d+\s*:\s*{", "{", content)
    start, end = cleaned.find("["), cleaned.rfind("]") + 1
    json_str = cleaned[start:end] if start != -1 and end != -1 else cleaned
    try:
        return json.loads(json_str)
    except json.JSONDecodeError:
        repls = []
        lines = cleaned.splitlines()
        for i, ln in enumerate(lines):
            if '"find"' in ln:
                fm = re.search(r'"find"\s*:\s*"([^"]*)"', ln)
                rm = None
                if fm:
                    for nxt in lines[i + 1:]:
                        m = re.search(r'"replace"\s*:\s*"([^"]*)"', nxt)
                        if m:
                            rm = m.group(1)
                            break
                if fm and rm:
                    repls.append({"find": fm.group(1), "replace": rm})
        return repls

# -----------------------------
# Modelaanroep
# -----------------------------
def get_replacements_from_model(groq_client: Groq, template_text: str, context_text: str):
    prompt = (
        "Gegeven TEMPLATE en CONTEXT, lever JSON-array met objecten {find, replace}."
        f"\n\nTEMPLATE:\n{template_text}\n\nCONTEXT:\n{context_text}"
    )
    try:
        resp = groq_client.chat.completions.create(
            model="llama-3.1-8b-instant",
            temperature=0.2,
            messages=[
                {"role": "system", "content": "Antwoord alleen JSON-array, geen extra tekst."},
                {"role": "user", "content": prompt},
            ],
        )
    except Exception as e:
        st.error(f"Fout bij model-aanroep: {e}")
        return []
    content = resp.choices[0].message.content
    repls = parse_groq_json_array(content)
    return [r for r in repls if r.get("find") and r.get("find") != r.get("replace")]

# -----------------------------
# Document bewerking
# -----------------------------
def apply_replacements_to_doc_and_bytes(doc_path: str, replacements: list[dict], include_changes_overview: bool = True) -> bytes:
    doc = docx.Document(doc_path)

    def repl(runs):
        if not runs:
            return
        txt = "".join(r.text for r in runs)
        for rp in replacements:
            txt = txt.replace(rp["find"], rp["replace"])
        runs[0].text = txt
        for r in runs[1:]:
            r.text = ""

    for p in doc.paragraphs:
        repl(p.runs)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    repl(p.runs)

    if include_changes_overview:
        try:
            p_br = doc.add_paragraph()
            run = p_br.add_run()
            run.add_break(WD_BREAK.PAGE)
        except Exception:
            pass
        try:
            doc.add_paragraph("Aangepaste onderdelen", style="Heading 1")
        except Exception:
            doc.add_paragraph("Aangepaste onderdelen")
        for rp in replacements:
            para = doc.add_paragraph(f"‚Ä¢ {rp['find']} ‚Üí {rp['replace']}")
            try:
                para.style = "List Bullet"
            except Exception:
                pass
        step_lines = steps.get_steps() if hasattr(steps, "get_steps") else []
        if step_lines:
            try:
                doc.add_paragraph("Uitgevoerde stappen", style="Heading 2")
            except Exception:
                doc.add_paragraph("Uitgevoerde stappen")
            for s in step_lines:
                doc.add_paragraph(f"‚Ä¢ {s}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# -----------------------------
# .docx veilig uitlezen
# -----------------------------
def _safe_read_docx_text(path: str) -> str:
    try:
        doc = docx.Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        return ""

# -----------------------------
# Streamlit UI - main
# -----------------------------
def run(show_nav: bool = True):
    try:
        from . import steps as _steps_mod
    except Exception:
        import importlib
        _steps_mod = importlib.import_module("tools.docex_tool.steps")

    steps_mod = _steps_mod
    steps_mod.clear_steps()

    st.set_page_config(page_title="DOCX Generator", layout="wide", initial_sidebar_state="expanded")
    st.markdown(
        """
        <style>
        .stButton>button, .stDownloadButton>button {font-size:18px; font-weight:bold; padding:0.6em 1.2em;}
        .big-header {font-size:2.5rem; font-weight:bold; margin-bottom:0.3em;}
        .section-header {font-size:1.75rem; font-weight:600; margin-top:1em; margin-bottom:0.5em;}
        .stTextArea textarea {font-family:monospace;}
        </style>
        """,
        unsafe_allow_html=True,
    )

    groq_client = get_groq_client()
    steps_mod.record_step("Groq client aangemaakt")

    st.markdown("<div class='big-header'>üöÄ Docex</div>", unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    tpl_path = None
    context = ""

    with col1:
        st.markdown("<div class='section-header'>üìÑ Template Upload</div>", unsafe_allow_html=True)
        tpl_file = st.file_uploader("Kies .docx template", type="docx", key="tpl")
        if tpl_file:
            tpl_path_dir = tempfile.mkdtemp()
            tpl_path = os.path.join(tpl_path_dir, "template.docx")
            with open(tpl_path, "wb") as f:
                f.write(tpl_file.getbuffer())
            tpl_text = _safe_read_docx_text(tpl_path)
            st.text_area("Template-inhoud", tpl_text, height=250, key="tpl_pre")
            steps_mod.record_step("Template ge√ºpload")

    with col2:
        st.markdown("<div class='section-header'>üìù Context Upload</div>", unsafe_allow_html=True)
        ctx_file = st.file_uploader("Kies .docx/.txt context", type=["docx", "txt"], key="ctx")
        if ctx_file:
            tmp_c = tempfile.mkdtemp()
            if ctx_file.type and ctx_file.type.endswith("document"):
                cpath = os.path.join(tmp_c, "context.docx")
                with open(cpath, "wb") as f:
                    f.write(ctx_file.getbuffer())
                context = _safe_read_docx_text(cpath)
                steps_mod.record_step("Context (.docx) ge√ºpload")
            else:
                context = ctx_file.read().decode("utf-8", errors="ignore")
                steps_mod.record_step("Context (.txt) ge√ºpload")
            st.text_area("Context-inhoud", context, height=250, key="ctx_pre")

    st.sidebar.markdown("### Uitgevoerde stappen")
    for s in steps_mod.get_steps():
        st.sidebar.markdown(f"- {s}")

    if tpl_path and context:
        st.markdown("---")
        if st.button("üéâ Genereer aangepast document"):
            steps_mod.record_step("Start generatie")
            tpl_text = _safe_read_docx_text(tpl_path)
            replacements = get_replacements_from_model(groq_client, tpl_text, context)
            steps_mod.record_step(f"Model returned {len(replacements)} replacement(s)")
            st.markdown("<div class='section-header'>‚ú® Aangepaste onderdelen</div>", unsafe_allow_html=True)
            if replacements:
                for rp in replacements:
                    st.write(f"‚Ä¢ **{rp['find']}** ‚Üí **{rp['replace']}**")
            else:
                st.info("Geen vervangingen gevonden.")
            out_bytes = apply_replacements_to_doc_and_bytes(tpl_path, replacements)
            steps_mod.record_step("Document gegenereerd (bytes klaar)")
            st.markdown("<div class='section-header'>üìù Volledige stappen</div>", unsafe_allow_html=True)
            for s in steps_mod.get_steps():
                st.write(f"- {s}")
            st.download_button(
                "‚¨áÔ∏è Download aangepast document",
                data=out_bytes,
                file_name="aangepast_template.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
    else:
        st.info("Upload eerst template en context om te starten.")

if __name__ == "__main__":
    run()
