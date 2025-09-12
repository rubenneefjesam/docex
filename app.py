import os
import tempfile
from docxtpl import DocxTemplate
import streamlit as st
import docx
from groq import Groq

# ─── Streamlit Page Config ───────────────────────────────────────────────
st.set_page_config(page_title="DOCX Generator", layout="wide")
st.title("📄 DOCX Generator met Templates")

# ─── Init Groq-client via st.secrets ──────────────────────────────────────
def get_groq_client():
    api_key = st.secrets.get("groq", {}).get("api_key", "").strip()
    if not api_key:
        st.sidebar.error("❌ Mist Groq-credentials! Voeg toe in `.streamlit/secrets.toml` onder [groq]")
        st.stop()
    try:
        client = Groq(api_key=api_key)
        st.sidebar.success("🔑 Verbonden met Groq API")
        return client
    except Exception as e:
        st.sidebar.error(f"❌ Fout bij verbinden met Groq API: {e}")
        st.stop()

groq_client = get_groq_client()

# ─── Document uitlezen ────────────────────────────────────────────────────
def read_docx(path: str) -> str:
    doc = docx.Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

# ─── Vul template via LLM ─────────────────────────────────────────────────
def enrich_and_fill(template_path: str, source_paths: list[str]) -> bytes:
    tpl_text = read_docx(template_path)
    ctx_parts = []
    for src in source_paths:
        text = read_docx(src)
        ctx_parts.append(f"=== CONTEXT FROM {os.path.basename(src)} ===\n{text}")
    context_text = "\n\n".join(ctx_parts)

    prompt = (
        "Werk de volgende template volledig bij op basis van de nieuwe context."
        f" Lever alleen de volledige bijgewerkte template-tekst terug, zonder toelichting of opsommingen."
        f"\n\n=== TEMPLATE ===\n{tpl_text}"
        f"\n\n{context_text}"
    )

    response = groq_client.chat.completions.create(
        model="llama-3.1-8b-instant",
        temperature=0.2,
        messages=[
            {"role": "system", "content": "Je bent een geavanceerde editor: geef alleen de aangepaste documenttekst terug zonder extra uitleg."},
            {"role": "user", "content": prompt}
        ]
    )
    updated = response.choices[0].message.content

    # Sla als DOCX
    out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc = docx.Document()
    for line in updated.split("\n"):
        doc.add_paragraph(line)
    doc.save(out.name)
    out.seek(0)
    return out.read()

# ─── Streamlit UI ────────────────────────────────────────────────────────
st.sidebar.header("Upload bestanden")
tpl_file = st.sidebar.file_uploader("Upload DOCX Template", type=["docx"])
src_files = st.sidebar.file_uploader("Upload Brondocumenten", type=["docx"], accept_multiple_files=True)

if tpl_file and src_files:
    tmp_dir = tempfile.mkdtemp()
    tpl_path = os.path.join(tmp_dir, "template.docx")
    with open(tpl_path, "wb") as f:
        f.write(tpl_file.getbuffer())

    src_paths = []
    for sf in src_files:
        p = os.path.join(tmp_dir, sf.name)
        with open(p, "wb") as o:
            o.write(sf.getbuffer())
        src_paths.append(p)

    st.subheader("Template Preview")
    st.write(read_docx(tpl_path)[:500] + ("…" if len(read_docx(tpl_path))>500 else ""))
    st.subheader("Context Preview")
    st.write(read_docx(src_paths[0])[:500] + ("…" if len(read_docx(src_paths[0]))>500 else ""))

    if st.button("Vul template aan met nieuwe/vervangende informatie"):
        st.info("Document wordt bijgewerkt…")
        try:
            result_bytes = enrich_and_fill(tpl_path, src_paths)
            st.download_button(
                "⬇️ Download bijgewerkt document",
                data=result_bytes,
                file_name="bijgewerkt_document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Bijwerken mislukt: {e}")
else:
    st.info("Upload zowel template als context om te starten.")