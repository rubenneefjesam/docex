"""
Use Case Analyzer Streamlit app

- Layout: two columns (links: korte invoer + knop; rechts: uitgewerkte use-case + download Word)
- Model: probeert eerst GROQ (GROQ_API_KEY), daarna OpenAI (OPENAI_API_KEY). Als geen API key beschikbaar gebruikt het een lokale template-fallback.

Dependencies:
  pip install streamlit python-docx openai groq

Run:
  streamlit run use_case_analyzer.py

"""

import os
import io
import textwrap
import streamlit as st
from docx import Document

# Optional model clients
try:
    from groq import Groq
    _has_groq = True
except Exception:
    _has_groq = False

try:
    import openai
    _has_openai = True
except Exception:
    _has_openai = False


# -----------------------------
# Model handling
# -----------------------------

def get_groq_client():
    if not _has_groq:
        return None
    api_key = os.environ.get("GROQ_API_KEY", "").strip()
    if not api_key:
        return None
    try:
        return Groq(api_key=api_key)
    except Exception:
        return None


def get_openai_api_key():
    return os.environ.get("OPENAI_API_KEY", "").strip()


def call_groq(prompt: str) -> str:
    client = get_groq_client()
    if not client:
        raise RuntimeError("Groq client niet beschikbaar of GROQ_API_KEY niet ingesteld")
    # Gebruik chat completion als beschikbaar
    try:
        resp = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            temperature=0.2,
            messages=[
                {"role": "system", "content": "Je bent een ervaren IT Business Analyst. Geef een nette, gestructureerde use-case in Nederlands."},
                {"role": "user", "content": prompt},
            ],
        )
        return resp.choices[0].message.content
    except Exception as e:
        raise RuntimeError(f"Groq-aanroep mislukt: {e}")


def call_openai_chat(prompt: str) -> str:
    key = get_openai_api_key()
    if not key or not _has_openai:
        raise RuntimeError("OpenAI niet beschikbaar of OPENAI_API_KEY niet ingesteld")
    openai.api_key = key
    try:
        # Gebruik ChatCompletion API (compatibiliteitsmodus)
        resp = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Je bent een ervaren IT Business Analyst. Geef een nette, gestructureerde use-case in Nederlands."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
            max_tokens=1200,
        )
        return resp.choices[0].message["content"] if "message" in resp.choices[0] else resp.choices[0].text
    except Exception as e:
        raise RuntimeError(f"OpenAI-aanroep mislukt: {e}")


def generate_usecase_text(short_input: str, include_acceptance: bool) -> str:
    """Genereer een gestructureerde use-case tekst.

    Eerst proberen we Groq, dan OpenAI, en als fallback gebruiken we een lokale generator.
    """
    # Prompt bouw
    prompt = textwrap.dedent(f"""
    Schrijf een volledige use-case in het Nederlands voor een IT Business Analyst.
    Gebruik duidelijke kopjes (Titel, Korte samenvatting, Scope, Actoren, Precondities,
    Hoofdscenario, Alternatieve flows, Acceptatiecriteria, Data & Integraties, Security/Compliance,
    Metrics, Next steps).

    Houd de taal zakelijk en bondig. Maak genummerde stappen voor het hoofdscenario.

    Input: {short_input}

    {'Voeg acceptatiecriteria toe.' if include_acceptance else 'Geen acceptatiecriteria nodig.'}
    """)

    # Probeer Groq
    try:
        if _has_groq and os.environ.get("GROQ_API_KEY", "").strip():
            return call_groq(prompt)
    except Exception:
        # fallback to next
        pass

    # Probeer OpenAI
    try:
        if _has_openai and get_openai_api_key():
            return call_openai_chat(prompt)
    except Exception:
        pass

    # Lokale fallback (geen API-key)
    return local_fallback_generator(short_input, include_acceptance)


def local_fallback_generator(short_input: str, include_acceptance: bool) -> str:
    # Eenvoudige deterministic generator die zonder model werkt
    title = generate_title_from_input(short_input)
    summary = f"Doel: {short_input.strip()}. Deze use-case beschrijft de gewenste functionaliteit en randvoorwaarden."
    scope = "In scope:\n- Basis upload en validatie\n- Boeking naar crediteuren\n\nOut of scope:\n- Geavanceerde OCR-tuning, bulk-verwerking"
    actors = "- Eindgebruiker (boekhouder)\n- Systeem (API)\n- Financials backoffice"
    pre = "- Gebruiker is ingelogd\n- Configuratie API-keys voor boekhoudsysteem aanwezig"
    main_flow = textwrap.dedent("""
    1. Gebruiker uploadt factuurbestand (PDF/PNG).
    2. Systeem valideert bestandsformaat en verplichte velden (bedrag, factuurnummer).
    3. Bij succesvolle validatie wordt de factuur geparseerd en mapping uitgevoerd.
    4. Systeem toont een preview en vraagt bevestiging van boeking.
    5. Na bevestiging wordt boeking aangemaakt in het crediteuren-systeem en een notificatie verstuurd.
    """)
    alt_flows = "- 2a: Validatie faalt → foutmelding en terug naar upload.\n- 4a: Gebruiker past velden aan → opnieuw valideren."
    acceptance = "- Facturen worden juist geparsed in 95% van de testset.\n- Succesvolle boeking resulteert in statuscode 200 van backend." if include_acceptance else ""
    data = "Belangrijke velden: factuurnummer, datum, totaalbedrag, leverancier, BTW-bedrag. Integratie: REST API naar ERP (POST /invoices)."
    security = "- Opslag van documenten in beveiligde bucket.\n- Persoonsgegevens volgens AVG verwerken.\n- Audit logging aanwezig."
    metrics = "- Time-to-book: < 5 minuten.\n- Foutpercentage validatie: < 5%."
    next_steps = "- Prototypen UI\n- Test dataset aanleveren\n- Integratie testen met test-ERP"

    parts = [
        f"Titel: {title}",
        "",
        f"Korte samenvatting:\n{summary}",
        "",
        f"Scope / Context:\n{scope}",
        "",
        f"Actoren:\n{actors}",
        "",
        f"Precondities:\n{pre}",
        "",
        f"Hoofdscenario:\n{main_flow}",
        "",
        f"Alternatieve flows:\n{alt_flows}",
    ]
    if include_acceptance:
        parts.append("")
        parts.append(f"Acceptatiecriteria:\n{acceptance}")
    parts.extend(["", f"Data & Integraties:\n{data}", "", f"Security / Compliance:\n{security}", "", f"Impact / Metrics:\n{metrics}", "", f"Next steps:\n{next_steps}"])

    return "\n".join(parts)


def generate_title_from_input(short_input: str) -> str:
    s = short_input.strip()
    if len(s) > 60:
        s = s[:57] + "..."
    # eenvoudige heuristiek: neem de eerste zin
    first = s.split('\n')[0]
    # vervang werkwoorden in infinitief context
    return f"Use-case: {first.capitalize()}"


# -----------------------------
# Word (.docx) generator
# -----------------------------

def create_docx_from_text(title: str, content: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    # Splits content in paragrafen op dubbele nieuwe regels
    for block in content.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        # If block looks like a heading (contains ':' as in 'Korte samenvatting:') render bold heading
        if ':' in block.splitlines()[0] and len(block.splitlines()[0]) < 60:
            h = block.splitlines()[0]
            rest = "\n".join(block.splitlines()[1:]).strip()
            doc.add_heading(h.rstrip(':'), level=2)
            if rest:
                for line in rest.split('\n'):
                    doc.add_paragraph(line)
        else:
            for line in block.split('\n'):
                doc.add_paragraph(line)

    bio = doc.core_properties
    bio.author = "Use Case Analyzer"
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# -----------------------------
# Streamlit UI
# -----------------------------

st.set_page_config(page_title="Use Case Analyzer", layout="wide")
st.title("📝 Use-case Analyzer voor IT Business Analysts")
st.write("Typ 1–4 regels en klik op 'Schrijf use-case uit' — rechts verschijnt de uitgewerkte use-case en een downloadknop (Word).")

col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("Invoer")
    user_input = st.text_area(
        label="Korte omschrijving (3–4 regels)",
        placeholder="Bijv. 'Gebruiker kan factuur uploaden; systeem valideert velden; bij succes wordt boeking gedaan naar crediteuren.'",
        height=140,
        max_chars=800,
        key="short_input",
    )
    char_count = len(user_input or "")
    st.caption(f"Tekens: {char_count}")
    include_acceptance = st.checkbox("Voeg acceptatiecriteria toe", value=True)
    generate_btn = st.button("Schrijf use-case uit")

    # Extra: toon beschikbare model keys
    st.markdown("**Model-config**")
    groq_present = bool(os.environ.get("GROQ_API_KEY", "").strip()) and _has_groq
    openai_present = bool(os.environ.get("OPENAI_API_KEY", "").strip()) and _has_openai
    st.write(f"Groq API: {'ja' if groq_present else 'nee'} | OpenAI API: {'ja' if openai_present else 'nee'}")
    st.markdown("---")
    st.write("Tip: als je geen API key zet, gebruikt de app een lokale fallback-template.")

with col2:
    st.subheader("Uitgewerkte use-case")
    output_placeholder = st.empty()

# Main action
if generate_btn:
    if not (user_input and user_input.strip()):
        st.warning("Vul een korte omschrijving in voordat je genereert.")
    else:
        with st.spinner("Bezig met genereren…"):
            try:
                out_text = generate_usecase_text(user_input, include_acceptance)
                st.session_state["latest_usecase"] = out_text
                st.session_state["latest_title"] = generate_title_from_input(user_input)
            except Exception as e:
                st.error(f"Fout bij genereren: {e}")
                out_text = None

        if out_text:
            # Toon in rechterkolom netjes: we hergebruiken output_placeholder
            with col2:
                output_placeholder.markdown(f"**{st.session_state.get('latest_title','Use-case')}**")
                # Render as markdown for readability
                output_placeholder.markdown("---")
                # We want to preserve line breaks
                output_placeholder.write(out_text)

            # Maak docx en bied download
            docx_bytes = create_docx_from_text(st.session_state.get("latest_title", "Use-case"), out_text)
            st.download_button(
                label="⬇️ Download use-case (Word)",
                data=docx_bytes,
                file_name="use_case.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            # Kopieer naar klembord (via kleine HTML/JS knop)
            copy_html = f"""
            <textarea id='uc' style='display:none;'>""" + out_text.replace("\n","\\n").replace('"','\"') + """</textarea>
            <button onclick="(async()=>{const t=document.getElementById('uc').value; await navigator.clipboard.writeText(t); this.innerText='Gekopieerd!';})()">Kopieer naar klembord</button>
            """
            st.components.v1.html(copy_html, height=50)

# Als er al een gegenereerde use-case in sessie staat — bied download ook aan
if st.session_state.get("latest_usecase") and not generate_btn:
    with col2:
        st.markdown(f"**{st.session_state.get('latest_title','Use-case')}**")
        st.markdown("---")
        st.write(st.session_state.get("latest_usecase"))
        docx_bytes = create_docx_from_text(st.session_state.get("latest_title", "Use-case"), st.session_state.get("latest_usecase"))
        st.download_button(
            label="⬇️ Download use-case (Word)",
            data=docx_bytes,
            file_name="use_case.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

# Footer / instructions
st.markdown("---")
st.caption("Notitie: Plaats je GROQ_API_KEY of OPENAI_API_KEY in je environment variables om een model te gebruiken. Zonder key wordt een lokale fallback-template gebruikt.")

# --- Compatibiliteitslaag voor importers die een symbol 'app' verwachten ---
def _compat_app_entrypoint():
    """
    Compat-entrypoint voor importers/registry die 'app' verwachten.
    Voor een Streamlit-script kunnen we hier geen server opstarten bij import,
    dus we geven een lichte descriptor / callable terug zonder side-effects.
    """
    # Je kunt hier ook metadata of een callable teruggeven als de registry
    # die later wil aanroepen. Pas aan naar behoefte.
    return {
        "type": "streamlit",
        "module": __name__,
        "run_cmd": "streamlit run " + __file__,
    }

# Exporteer 'app' zodat 'from ... import app' werkt.
# We exporteren de callable zelf (dus importeren geeft een callable terug).
app = _compat_app_entrypoint
