import os
import io
import tempfile
import streamlit as st
from groq import Groq
import docx
from docx import Document
from docx.enum.text import WD_BREAK
import csv
import io as _io

# ----------------------------------
# Helper: Groq client ophalen
# ----------------------------------

def get_groq_client():
    api_key = os.environ.get("GROQ_API_KEY", "").strip()
    if api_key:
        try:
            return Groq(api_key=api_key)
        except Exception as e:
            st.sidebar.error(f"❌ Fout bij verbinden met Groq API: {e}")
            st.stop()
    try:
        api_key = st.secrets.get("groq", {}).get("api_key", "").strip()
    except Exception:
        api_key = ""
    if not api_key:
        st.sidebar.error(
            "❌ Groq API key niet gevonden. Zet GROQ_API_KEY als env var of gebruik .streamlit/secrets.toml"
        )
        st.stop()
    try:
        return Groq(api_key=api_key)
    except Exception as e:
        st.sidebar.error(f"❌ Fout bij verbinden met Groq API: {e}")
        st.stop()

# ----------------------------------
# Helper: plain text uit docx
# ----------------------------------

def _read_docx_text(path: str) -> str:
    try:
        doc = docx.Document(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        return ""

# ----------------------------------
# Generate beheersplan voor één risico
# ----------------------------------

def generate_plan_for_risk(groq_client, examples: str, record: dict) -> str:
    # Maak context met risico, oorzaak en gevolg
    context = (
        f"Risico: {record['risico']}\n"
        f"Oorzaak: {record.get('oorzaak', '')}\n"
        f"Gevolg: {record.get('gevolg', '')}"
    )
    prompt = (
        f"Gegeven voorbeelden van bestaande beheersplannen:\n{examples}\n"
        f"Schrijf een nieuw beheersplan (maximaal twee alinea's) voor het volgende risico, inclusief oorzaak en gevolg:\n{context}\n"
        "Antwoord in platte tekst zonder extra uitleg."
    )
    resp = groq_client.chat.completions.create(
        model="llama-3.1-8b-instant",
        temperature=0.3,
        messages=[
            {"role": "system", "content": "Je bent een risico-expert. Geef alleen het beheersplan."},
            {"role": "user", "content": prompt},
        ],
    )
    return resp.choices[0].message.content.strip()

# ----------------------------------
# Parse risico-bestand
# ----------------------------------

def parse_risks(text: str) -> list[dict]:
    lines = [l for l in text.splitlines() if l.strip()]
    records = []
    # Probeer CSV met header
    try:
        reader = csv.DictReader(_io.StringIO(text), delimiter=',')
        if {'risico','oorzaak','gevolg'}.issubset(reader.fieldnames):
            for row in reader:
                records.append({k: row[k].strip() for k in ['risico','oorzaak','gevolg']})
            return records
    except Exception:
        pass
    # Fallback: iedere 3 regels vormt één record
    for i in range(0, len(lines), 3):
        chunk = lines[i:i+3]
        if len(chunk) >= 1:
            rec = {'risico': chunk[0].strip()}
            if len(chunk) > 1 and ':' in chunk[1]:
                rec['oorzaak'] = chunk[1].split(':',1)[1].strip()
            if len(chunk) > 2 and ':' in chunk[2]:
                rec['gevolg'] = chunk[2].split(':',1)[1].strip()
            records.append(rec)
    return records

# ----------------------------------
# Streamlit UI - hoofd
# ----------------------------------

def run():
    st.set_page_config(page_title="Risk Plan Generator", layout="wide")
    st.title("Risico Beheersplan Generator 🐯")

    groq_client = get_groq_client()

    # Uploaders in twee kolommen
    col1, col2 = st.columns(2)
    examples_text = ""
    risks_text = ""
    with col1:
        st.subheader("📄 Upload oude beheersplannen (docx)")
        old_files = st.file_uploader(
            "Kies één of meerdere .docx bestanden", type="docx", accept_multiple_files=True
        )
        if old_files:
            parts = []
            for f in old_files:
                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                tmp.write(f.getbuffer())
                tmp.flush()
                parts.append(_read_docx_text(tmp.name))
            examples_text = "\n---\n".join(parts)
    with col2:
        st.subheader("📝 Upload risico's (txt/csv)")
        risk_file = st.file_uploader(
            "Kies een .txt of .csv bestand met kolommen: risico, oorzaak, gevolg", type=["txt","csv"]
        )
        if risk_file:
            risks_text = risk_file.read().decode("utf-8", errors="ignore")

    if st.button("🎉 Genereer beheersplannen"):  
        if not examples_text:
            st.error("Upload eerst bestaande beheersplannen in kolom 1.")
            return
        if not risks_text:
            st.error("Upload eerst risico's in kolom 2.")
            return

        records = parse_risks(risks_text)
        if not records:
            st.error("Geen geldige risico-records gevonden. Controleer je bestandstructuur.")
            return

        st.info("Genereren kan even duren...")
        plans = []
        for rec in records:
            plan = generate_plan_for_risk(groq_client, examples_text, rec)
            plans.append((rec, plan))

        # Toon output
        st.markdown("## 📑 Resultaat")
        for rec, plan in plans:
            st.markdown(f"**Risico:** {rec['risico']}")
            if rec.get('oorzaak'):
                st.markdown(f"- Oorzaak: {rec['oorzaak']}")
            if rec.get('gevolg'):
                st.markdown(f"- Gevolg: {rec['gevolg']}")
            st.write(plan)
            st.markdown("---")

        # Exporteer naar Word
        buffer = io.BytesIO()
        doc = Document()
        for rec, plan in plans:
            doc.add_paragraph(f"Risico: {rec['risico']}", style="Heading 2")
            if rec.get('oorzaak'):
                doc.add_paragraph(f"Oorzaak: {rec['oorzaak']}")
            if rec.get('gevolg'):
                doc.add_paragraph(f"Gevolg: {rec['gevolg']}")
            for para in plan.split("\n\n"):
                doc.add_paragraph(para)
            doc.add_page_break()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="⬇️ Exporteer naar Word",
            data=buffer.read(),
            file_name="beheersplannen.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    run()
