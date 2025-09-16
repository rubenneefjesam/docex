import os
import io
import streamlit as st
from groq import Groq
from docx import Document

# ─── Configuratie ────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Meeting2Document",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── Groq Client Initialisatie ────────────────────────────────────────────────
@st.cache_resource
def init_groq_client():
    # Eerst uit omgevingsvariabelen, dan uit Streamlit-secrets
    key = os.getenv("GROQ_API_KEY", "").strip() \
          or st.secrets.get("groq", {}).get("api_key", "").strip()
    if not key:
        st.warning("⚠️ Geen Groq-API-key gevonden; transcriptie en AI-generatie werken niet.")
        return None
    try:
        return Groq(api_key=key)
    except Exception:
        st.error("❌ Groq-API-key ongeldig.")
        return None

client = init_groq_client()

# ─── Helper-functies ──────────────────────────────────────────────────────────
def transcribe_audio(audio_file):
    """Transcribeer audio met Groq Whisper."""
    audio_data = audio_file.read()
    st.audio(audio_data, format=audio_file.type)
    with st.spinner("Transcriberen audio…"):
        res = client.audio.transcriptions.create(
            model="whisper-large-v3",
            file=(audio_file.name, audio_data),
        )
    return res.text

def generate_minutes(transcript_text: str) -> str:
    """Genereer vergaderverslag met Groq chat completions."""
    system_msg = "Je bent een assistent voor het maken van vergaderverslagen."
    user_msg = (
        "Maak op basis van de volgende meeting transcriptie een gestructureerd verslag. "
        "Gebruik de kopjes: Samenvatting, Aanwezigen, Actiepunten, Besluiten.\n\n"
        f"Transcriptie:\n{transcript_text}"
    )
    with st.spinner("Bezig met AI-generatie…"):
        resp = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            temperature=0.3,
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user",   "content": user_msg},
            ],
        )
    return resp.choices[0].message.content.strip()

def build_docx(minute_text: str) -> io.BytesIO:
    """Bouwt een .docx bestand van de gegenereerde notulen."""
    buffer = io.BytesIO()
    doc = Document()

    # Verwerk regels: splits op nieuwe regels
    for line in minute_text.split("\n"):
        if line.startswith("Samenvatting"):
            doc.add_heading("Samenvatting", level=2)
            content = line.partition(":")[2].strip()
            doc.add_paragraph(content)
        elif line.startswith("Aanwezigen"):
            doc.add_heading("Aanwezigen", level=2)
            content = line.partition(":")[2].strip()
            doc.add_paragraph(content)
        elif line.startswith("Actiepunten"):
            doc.add_heading("Actiepunten", level=2)
            content = line.partition(":")[2].strip()
            doc.add_paragraph(content)
        elif line.startswith("Besluiten"):
            doc.add_heading("Besluiten", level=2)
            content = line.partition(":")[2].strip()
            doc.add_paragraph(content)
        else:
            # reguliere alinea
            doc.add_paragraph(line)

    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ─── Navigatie ────────────────────────────────────────────────────────────────
tabs = st.tabs(["Home", "Upload & Generate", "Over"])

# ─── Tab: Home ───────────────────────────────────────────────────────────────
with tabs[0]:
    st.title("✨ Welkom bij Meeting2Document")
    st.markdown(
        """
        Deze tool zet je **meeting transcript** of **notities** om in een professioneel **vergaderverslag** (.docx).

        **Hoe werkt het?**
        1. Ga naar **Upload & Generate**  
        2. Upload je transcript (.txt) of audio (.wav, .mp3)  
        3. Klik op **Genereer verslag**  
        4. Download je .docx-verslag  
        """
    )

# ─── Tab: Upload & Generate ─────────────────────────────────────────────────
with tabs[1]:
    st.title("📂 Upload je notities of audio")

    transcript_text = st.session_state.get("transcript_text", "")

    # Audio upload en transcriptie
    st.subheader("🎤 Upload audio (.wav, .mp3)")
    audio_file = st.file_uploader("Kies audiobestand", type=["wav", "mp3"])
    if audio_file and client:
        try:
            transcript_text = transcribe_audio(audio_file)
            st.success("Transcriptie audio voltooid ✅")
            st.session_state["transcript_text"] = transcript_text
        except Exception as e:
            st.error(f"Transcriptie mislukt: {e}")

    # Fallback: tekstbestand upload
    st.subheader("🗒️ Upload meeting transcript (.txt)")
    text_file = st.file_uploader("Kies je meeting transcript", type=["txt"])
    if text_file:
        transcript_text = text_file.read().decode("utf-8", errors="ignore")
        st.session_state["transcript_text"] = transcript_text

    # Preview van transcript
    if transcript_text:
        st.subheader("📄 Transcript Preview")
        st.text_area("", transcript_text, height=200)

    # Genereer knop
    if st.button("✅ Genereer vergaderverslag"):
        if not transcript_text:
            st.error("Upload eerst audio of transcriptie.")
        elif not client:
            st.error("Geen geldige Groq-client. Controleer je API-key.")
        else:
            minutes = generate_minutes(transcript_text)
            st.subheader("🧾 Vergaderverslag")
            st.write(minutes)

            # Bouw en bied .docx download aan
            doc_buffer = build_docx(minutes)
            st.download_button(
                label="⬇️ Download verslag (.docx)",
                data=doc_buffer,
                file_name="vergaderverslag.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

# ─── Tab: Over ───────────────────────────────────────────────────────────────
with tabs[2]:
    st.title("ℹ️ Over deze tool")
    st.markdown(
        """
        **Meeting2Document** zet meetingnotities om in een professioneel verslag.

        **Features:**
        - Automatische transcriptie van audio (Whisper v3)
        - Samenvatting van de meeting
        - Overzicht van aanwezigen
        - Lijst met actiepunten en besluiten
        - Downloadbaar als Word-document (.docx)
        """
    )

# ─── Entrypoint ───────────────────────────────────────────────────────────────
if __name__ == "__main__":
    pass  # Streamlit voert automatisch de module uit
