# use_case_analyzer.py
# Streamlit-app voor inladen, genereren en exporteren van use-case templates.
# Verbeterde, robuustere versie — Groq wordt altijd gebruikt (geen checkbox).
# Extra guards toegevoegd tegen IndexError bij DOCX-export, veiliger JSON-extractie
# en robuustere parsing van Groq-responses.

import os
import re
import json
import textwrap
import io
from typing import Dict, Tuple, Optional, Any
import streamlit as st

# Optionele imports
try:
    from groq import Groq
    _HAS_GROQ = True
except Exception:
    Groq = None
    _HAS_GROQ = False

try:
    from docx import Document
    _HAS_DOCX = True
except Exception:
    Document = None
    _HAS_DOCX = False

# Regex voor placeholders {{ key }}
_PLACEHOLDER_RE = re.compile(r"{{\s*([a-zA-Z0-9_]+)\s*}}")


# ---------------------------
# Helper / Core functies
# ---------------------------

def load_templates(templates_dir: str) -> Dict[str, str]:
    """Laad alle templates uit een directory en geef terug als dict name -> content."""
    if not os.path.isdir(templates_dir):
        return {}
    templates = {}
    for fn in sorted(os.listdir(templates_dir)):
        path = os.path.join(templates_dir, fn)
        if fn.startswith('.') or not os.path.isfile(path):
            continue
        key = os.path.splitext(fn)[0]
        with open(path, encoding='utf-8') as f:
            templates[key] = f.read()
    return templates


def extract_placeholders(template_str: str) -> Tuple[str, ...]:
    """Vind placeholders en behoud de volgorde zonder duplicaten."""
    found = _PLACEHOLDER_RE.findall(template_str)
    seen, ordered = set(), []
    for ph in found:
        if ph not in seen:
            seen.add(ph)
            ordered.append(ph)
    return tuple(ordered)


def render_template(template_str: str, context: Dict[str, str]) -> str:
    """Rendert template door placeholders te vervangen met waarden uit context."""
    def repl(m: re.Match) -> str:
        return context.get(m.group(1), '')
    return _PLACEHOLDER_RE.sub(repl, template_str)


def _get_groq_client() -> Optional[Groq]:
    """
    Maak en retourneer een Groq-client wanneer mogelijk.
    Controleert eerst environment, daarna st.secrets.
    """
    if not _HAS_GROQ:
        return None
    api_key = os.environ.get('GROQ_API_KEY', '').strip()
    try:
        # st.secrets kan in sommige test-omgevingen ontbreken; daarom try/except
        api_key = api_key or st.secrets.get('groq', {}).get('api_key', '').strip()
    except Exception:
        pass
    if not api_key:
        return None
    try:
        return Groq(api_key=api_key)
    except Exception:
        # Client aanmaken faalde (bijv. verkeerde SDK-versie). Retourneer None.
        return None


def _extract_first_json_object(s: str) -> Optional[str]:
    """
    Probeer uit string s het eerste json-object (balancerende {}) te extraheren.
    Retourneer het substring of None als niet gevonden.
    """
    start = s.find('{')
    if start == -1:
        return None
    depth = 0
    for i in range(start, len(s)):
        ch = s[i]
        if ch == '{':
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0:
                return s[start:i+1]
    return None


def _resp_to_text(resp: Any) -> str:
    """
    Converteer een Groq-response object naar bruikbare tekst.
    Probeert meerdere toegangen (attribuut- en dict-stijl) en valt terug op str(resp).
    """
    try:
        # Veel SDKs hebben resp.choices[...] maar de shape kan variëren.
        if hasattr(resp, 'choices'):
            choices = getattr(resp, 'choices')
            if choices and len(choices) > 0:
                choice = choices[0]
                # probeer message.content
                msg = getattr(choice, 'message', None)
                if msg is not None:
                    content = getattr(msg, 'content', None)
                    if content:
                        return content
                # probeer choice.text of choice.content
                content = getattr(choice, 'text', None) or getattr(choice, 'content', None)
                if content:
                    return content
        # Als resp een dict-achtige mapping is:
        if isinstance(resp, dict):
            choices = resp.get('choices') or []
            if choices:
                c0 = choices[0]
                if isinstance(c0, dict):
                    # zoek veelvoorkomende plekken
                    msg = c0.get('message') or {}
                    content = msg.get('content') if isinstance(msg, dict) else None
                    if content:
                        return content
                    # fallback
                    for key in ('text', 'content'):
                        if key in c0 and c0[key]:
                            return c0[key]
        # uiteindelijk fallback naar stringrepresentatie
        return str(resp)
    except Exception:
        return str(resp)


def call_groq(prompt: str, model: str = 'llama-3.1-8b-instant', temperature: float = 0.2) -> str:
    """
    Roept Groq API aan en retourneert de ruwe tekst (wat het model produceert).
    Gooit een duidelijke RuntimeError als de client niet beschikbaar is of de call faalt.
    """
    client = _get_groq_client()
    if not client:
        raise RuntimeError('Groq client niet beschikbaar. Controleer GROQ_API_KEY in omgeving of st.secrets.')
    try:
        resp = client.chat.completions.create(
            model=model,
            temperature=temperature,
            messages=[
                {'role': 'system', 'content': 'Je bent een ervaren IT Business Analyst.'},
                {'role': 'user', 'content': prompt}
            ]
        )
    except Exception as e:
        # Maak foutmelding duidelijk inclusief (gekorte) repr van exception
        raise RuntimeError(f'Fout bij aanroepen van Groq API: {e}')
    # Converteer response naar bruikbare tekst op een robuuste manier
    text = _resp_to_text(resp)
    return text


def local_generate_mapping(short_input: str, placeholders: Tuple[str, ...]) -> Dict[str, str]:
    """Eenvoudige fallback mapping voor ontbrekende placeholders."""
    first = short_input.strip().splitlines()[0] if short_input else 'Onbekende story'
    mapping = {}
    for ph in placeholders:
        low = ph.lower()
        if low == 'title':
            mapping[ph] = first
        elif low == 'role':
            mapping[ph] = 'Als gebruiker'
        elif low == 'goal':
            mapping[ph] = f'Wil ik {first.lower()} zodat ik waarde kan behalen.'
        elif low in ('steps', 'main_flow'):
            mapping[ph] = textwrap.dedent(
                '1. Gebruiker opent de feature.\n'
                '2. Gebruiker voert input in.\n'
                '3. Systeem valideert en bevestigt.\n'
                '4. Actie voltooid.'
            )
        else:
            mapping[ph] = first
    return mapping


def generate_for_template(short_input: str, template_name: str, templates_dir: str, use_groq: bool = True) -> Dict[str, str]:
    """
    Genereer waarden voor placeholders van een template.
    Gebruik Groq altijd wanneer use_groq=True (en beschikbaar). Valt terug op local_generate_mapping per missende velden.
    """
    templates = load_templates(templates_dir)
    tpl = templates.get(template_name)
    if tpl is None:
        raise FileNotFoundError(f"Template '{template_name}' niet gevonden.")
    placeholders = extract_placeholders(tpl)
    if not placeholders:
        return {}

    data: Dict[str, str] = {}

    if use_groq:
        try:
            prompt = textwrap.dedent(f"""
                Geef **alleen** geldig JSON met velden: {', '.join(placeholders)}
                Template: {template_name}
                Omschrijving: {short_input}
            """)
            raw = call_groq(prompt)
            # Probeer direct te parsen; als dat faalt, probeer de eerste JSON-object extractie
            parsed = None
            try:
                parsed = json.loads(raw)
            except Exception:
                # probeer eerste {} block te vinden en te parsen
                blob = _extract_first_json_object(raw)
                if blob:
                    try:
                        parsed = json.loads(blob)
                    except Exception:
                        parsed = None
            if isinstance(parsed, dict):
                data = {k: str(parsed.get(k, '')) for k in placeholders}
            else:
                # parsed niet bruikbaar: leave data leeg -> fallback verder
                # bewaar raw response in session_state voor debug (zichtbaar in UI)
                st.session_state['last_groq_raw'] = raw
        except Exception as e:
            # Zet debug info en laat fallback downstream de ontbrekende velden invullen
            st.session_state['last_groq_error'] = str(e)
            st.session_state['last_groq_raw'] = st.session_state.get('last_groq_raw', '')
            data = {}
    # Vul ontbrekende placeholders met local mapping
    missing = [ph for ph in placeholders if not data.get(ph)]
    if missing:
        data.update(local_generate_mapping(short_input, tuple(missing)))
    return data


def render_and_export(short_input: str, template_name: str, templates_dir: str, use_groq: bool = True) -> Tuple[str, bytes]:
    """
    Render de template met de gegenereerde data en exporteer naar docx (indien python-docx geïnstalleerd) of txt.
    Bevat guards tegen lege paragrafen om IndexError te voorkomen.
    """
    templates = load_templates(templates_dir)
    if template_name not in templates:
        raise FileNotFoundError(f"Template '{template_name}' niet gevonden.")
    data_map = generate_for_template(short_input, template_name, templates_dir, use_groq)
    tpl = templates[template_name]
    content = render_template(tpl, data_map)

    if _HAS_DOCX:
        doc = Document()
        # veilige heading (geen None)
        doc.add_heading(data_map.get('title') or template_name, level=1)
        # split op dubbele newlines, maar guard voor lege blokken
        for block in content.split('\n\n'):
            lines = [ln for ln in block.splitlines() if ln.strip() != '']
            if not lines:
                # sla lege blocks over
                continue
            first_line = lines[0]
            # als eerste regel een kop lijkt (bevat ':' aan het einde of in het midden)
            if ':' in first_line:
                # heading zonder trailing ':' 
                doc.add_heading(first_line.rstrip(':'), level=2)
                body_lines = lines[1:]
            else:
                body_lines = lines
            for ln in body_lines:
                if ln.strip():
                    doc.add_paragraph(ln)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        data = buf.read()
    else:
        data = content.encode('utf-8')
    return content, data


# ---------------------------
# Streamlit UI: main page layout
# ---------------------------

def app():
    st.set_page_config(page_title='Use-case Analyzer', layout='wide')
    st.title('📋 Use-case Analyzer')
    st.markdown(
        '## Werkwijze  \n'
        '1. Selecteer een template  \n'
        '2. Vul een korte omschrijving in  \n'
        '3. Klik op **Genereer use-case**  \n'
        '4. Bekijk en download het resultaat'
    )

    # Basis checks en template load
    here = os.path.dirname(__file__)
    templates_dir = os.path.join(here, 'templates')
    templates = load_templates(templates_dir)
    if not templates:
        st.error("Geen templates in 'templates/'. Plaats .txt/.md templates in de templates-map.")
        return

    # Forceer gebruik van Groq: controleer of Groq-client daadwerkelijk beschikbaar is
    groq_client = _get_groq_client()
    if groq_client is None:
        st.error(
            "Groq client niet beschikbaar. Deze applicatie vereist Groq.\n"
            "Zorg dat de Groq SDK geïnstalleerd is en dat je een geldige API key hebt "
            "in de omgeving (GROQ_API_KEY) of in st.secrets['groq']['api_key']."
        )
        # Toon debug hints indien aanwezig
        return

    col1, col2 = st.columns([1, 2])
    with col1:
        choice = st.selectbox('Kies template', list(templates.keys()))
        desc = st.text_area('Korte omschrijving', height=150)
        # Geen checkbox meer: Groq wordt altijd gebruikt
        st.markdown("**Generatie methode:** Groq LLM (verplicht).")

        gen = st.button('Genereer use-case')

    if gen:
        try:
            # Gebruik expliciet use_groq=True
            content, filedata = render_and_export(desc, choice, templates_dir, use_groq=True)
            st.subheader('Resultaat')
            st.code(content, language='markdown')
            ext = 'docx' if _HAS_DOCX else 'txt'
            mime = (
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                if ext == 'docx' else 'text/plain'
            )
            st.download_button(
                f'Download .{ext}', filedata,
                file_name=f'{choice}.{ext}', mime=mime
            )

            # Toon optioneel debug info van laatste Groq call (handig bij fouten)
            if 'last_groq_error' in st.session_state:
                with st.expander('Groq error (debug)'):
                    st.write(st.session_state.get('last_groq_error'))
                    st.write('Raw Groq response (indien aanwezig):')
                    st.code(st.session_state.get('last_groq_raw', ''), language='json')
            elif 'last_groq_raw' in st.session_state:
                with st.expander('Laatste Groq raw response'):
                    st.code(st.session_state.get('last_groq_raw', ''), language='json')

        except Exception as e:
            # Toon een concrete foutmelding en, indien beschikbaar, extra debug info
            st.error(f'Fout tijdens generatie/export: {e}')
            if 'last_groq_raw' in st.session_state:
                with st.expander('Raw Groq response (debug)'):
                    st.code(st.session_state.get('last_groq_raw', ''), language='json')
            if 'last_groq_error' in st.session_state:
                with st.expander('Groq error (debug)'):
                    st.write(st.session_state.get('last_groq_error'))
