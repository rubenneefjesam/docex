"""
story_writer.py

Helpers om 'story' (en andere) templates te vullen:
- load_templates(templates_dir)
- generate_for_template(short_input, template_name, templates_dir, use_groq=True)
- render_template(template_str, context)
- create_docx_from_text(title, content)

Werkt zonder Groq of python-docx (heeft fallbacks).
"""

import os
import re
import json
import textwrap
import io
from typing import Dict, Tuple, Optional

# Optioneel: Groq importeren (veilig bij ontbreken)
try:
    from groq import Groq
    _HAS_GROQ = True
except Exception:
    Groq = None
    _HAS_GROQ = False

# Optioneel: python-docx (voor download). Niet fataal als afwezig.
try:
    from docx import Document
    _HAS_DOCX = True
except Exception:
    Document = None
    _HAS_DOCX = False


# ----------------------------
# Template I/O & parsing
# ----------------------------

def load_templates(templates_dir: str) -> Dict[str, str]:
    """
    Lees alle templates in templates_dir en geef dict terug: naam -> inhoud.
    Naam is filename zonder extensie. Ignore hidden files.
    """
    templates = {}
    if not os.path.isdir(templates_dir):
        return templates
    for fn in sorted(os.listdir(templates_dir)):
        if fn.startswith("."):
            continue
        path = os.path.join(templates_dir, fn)
        if os.path.isfile(path):
            name = os.path.splitext(fn)[0]
            with open(path, encoding="utf-8") as f:
                templates[name] = f.read()
    return templates


_placeholder_re = re.compile(r"{{\s*([a-zA-Z0-9_]+)\s*}}")

def extract_placeholders(template_str: str) -> Tuple[str, ...]:
    """Retourneer tuple van unieke placeholder namen in volgorde van voorkomen."""
    found = _placeholder_re.findall(template_str)
    # preserve order, unique
    seen = set()
    order = []
    for k in found:
        if k not in seen:
            seen.add(k)
            order.append(k)
    return tuple(order)


def render_template(template_str: str, context: Dict[str, str]) -> str:
    """
    Vervang {{key}} in template_str met context[key] (of lege string).
    Houdt structurele newlines intact.
    """
    def repl(m):
        key = m.group(1)
        val = context.get(key, "")
        # als lijst-achtige inhoud (n lijnen) zorgen we dat de indentatie blijft
        return val or ""
    return _placeholder_re.sub(repl, template_str)


# ----------------------------
# Groq helpers
# ----------------------------

def get_groq_client() -> Optional[object]:
    """Probeer Groq-client te initialiseren op basis van env / secrets."""
    if not _HAS_GROQ:
        return None
    key = os.environ.get("GROQ_API_KEY", "").strip()
    # probeer fallback (bijv. streamlit secrets) zonder import dependency
    try:
        import streamlit as st  # type: ignore
        key = key or st.secrets.get("groq", {}).get("api_key", "").strip()
    except Exception:
        pass
    if not key:
        return None
    try:
        return Groq(api_key=key)
    except Exception:
        return None


def call_groq(prompt: str, model: str = "llama-3.1-8b-instant", temperature: float = 0.2) -> str:
    """
    Roep Groq aan en geef ruwe tekst terug. Werpt RuntimeError als client niet beschikbaar is.
    """
    client = get_groq_client()
    if not client:
        raise RuntimeError("Groq client niet beschikbaar of GROQ_API_KEY niet gezet")
    resp = client.chat.completions.create(
        model=model,
        temperature=temperature,
        messages=[
            {"role": "system", "content": "Je bent een ervaren IT Business Analyst. Antwoord precies zoals gevraagd."},
            {"role": "user", "content": prompt},
        ],
    )
    # veilige toegang tot response
    try:
        return resp.choices[0].message.content
    except Exception:
        # fallback: probeer string-conversie
        return str(resp)


# ----------------------------
# Fallback generator (lokaal)
# ----------------------------

def local_generate_mapping(short_input: str, placeholders: Tuple[str, ...]) -> Dict[str, str]:
    """
    Maak een eenvoudige mapping voor elke placeholder als Groq niet beschikbaar is.
    Logica is heuristisch: zoekt role/goal/steps/acceptance_criteria etc.
    """
    short = (short_input or "").strip()
    first_line = (short.splitlines()[0] if short else "Onbekende story")[:120]

    # basiscontent
    mapping = {}
    for key in placeholders:
        low = key.lower()
        if low in ("title", "naam", "naam_title"):
            mapping[key] = first_line
        elif low in ("role", "actor", "actoren"):
            mapping[key] = "Als gebruiker"
        elif low in ("goal", "doel", "want"):
            mapping[key] = f"Wil ik {first_line.lower()} zodat ik waarde kan behalen."
        elif low in ("reason", "reden"):
            mapping[key] = "Om tijd te besparen en fouten te verminderen."
        elif low in ("steps", "main_flow", "hoofdscenario"):
            # voorbeeldstappen; render als genummerde lijst (multiline)
            mapping[key] = textwrap.dedent("""\
                1. Gebruiker opent de functie.
                2. Gebruiker voert benodigde gegevens in.
                3. Systeem valideert invoer en geeft bevestiging.
                4. Actie wordt uitgevoerd en resultaat opgeslagen.""")
        elif low in ("acceptance_criteria", "acceptatiecriteria"):
            mapping[key] = "- E2E flow werkt\n- Validatie condities zijn voldaan"
        elif low in ("preconditions", "precondities"):
            mapping[key] = "- Gebruiker is ingelogd\n- Relevante data bestaat"
        else:
            # generieke fallback: echo kort input als context
            mapping[key] = f"{first_line} — aanvullende info: {short[:240]}"
    return mapping


# ----------------------------
# Main generator: probeer Groq -> fallback
# ----------------------------

def _build_groq_prompt_for_placeholders(short_input: str, placeholders: Tuple[str, ...], template_name: str) -> str:
    """
    Bouw prompt dat Groq vraagt om strikt geldige JSON met de benodigde keys.
    """
    ph_list = ", ".join(placeholders) if placeholders else ""
    prompt = textwrap.dedent(f"""
    Je krijgt een korte omschrijving van een user story / feature / epic en een template met placeholders.
    Geef uitsluitend **geldig JSON** terug (geen extra tekst) met de volgende velden: {ph_list}
    - Elke waarde moet een string zijn.
    - Voor velden als 'steps' of 'main_flow' geef meerdere regels (nieuw-regel gescheiden).
    - Houd het Nederlands aan.

    Template-type: {template_name}
    Korte omschrijving: {short_input}
    """)
    return prompt


def generate_for_template(short_input: str, template_name: str, templates_dir: str, use_groq: bool = True) -> Dict[str, str]:
    """
    Genereer mapping voor placeholders van template_name.
    Retourneert dict met keys=placeholders en values=gegenereerde tekst.
    Indien Groq faalt of niet beschikbaar: lokale fallback.
    """
    templates = load_templates(templates_dir)
    tpl = templates.get(template_name)
    if tpl is None:
        raise FileNotFoundError(f"Template '{template_name}' niet gevonden in {templates_dir}. Beschikbare: {list(templates.keys())}")

    placeholders = extract_placeholders(tpl)
    if not placeholders:
        # geen placeholders -> return lege mapping
        return {}

    # Probeer Groq
    if use_groq and _HAS_GROQ and get_groq_client():
        prompt = _build_groq_prompt_for_placeholders(short_input, placeholders, template_name)
        try:
            resp_text = call_groq(prompt)
            # probeer substring JSON-extractie
            try:
                start = resp_text.index("{")
                end = resp_text.rindex("}")
                json_text = resp_text[start:end+1]
                data = json.loads(json_text)
            except Exception:
                # fallback: als resp_text looks like kv lines, probeer heuristisch
                data = {}
                try:
                    # probeer rechtstreeks json.loads als resp_text is proper json
                    data = json.loads(resp_text)
                except Exception:
                    raise ValueError("Kon JSON niet parsen uit Groq-response")
            # zorg dat alle placeholders aanwezig zijn (anders toevoegen via fallback)
            for k in placeholders:
                if k not in data:
                    data[k] = ""
            # convert all values to strings
            for k in list(data.keys()):
                if data[k] is None:
                    data[k] = ""
                elif not isinstance(data[k], str):
                    # lijsten -> nieuwe lijnen
                    if isinstance(data[k], (list, tuple)):
                        data[k] = "\n".join(str(x) for x in data[k])
                    else:
                        data[k] = str(data[k])
            # vul ontbrekende keys met local fallback
            missing = [k for k in placeholders if not data.get(k)]
            if missing:
                local = local_generate_mapping(short_input, tuple(missing))
                for k in missing:
                    data[k] = local.get(k, "")
            return {k: data.get(k, "") for k in placeholders}
        except Exception as e:
            # log/print kan hier handig zijn tijdens ontwikkeling
            # fallback naar lokaal
            return local_generate_mapping(short_input, placeholders)
    else:
        # geen Groq beschikbaar -> local fallback
        return local_generate_mapping(short_input, placeholders)


# ----------------------------
# Document creation (docx fallback)
# ----------------------------

def create_docx_from_text(title: str, content: str) -> bytes:
    """
    Maak een .docx bestand (bytes). Als python-docx ontbreekt, return plain text bytes (utf-8).
    """
    if not _HAS_DOCX:
        return content.encode("utf-8")
    doc = Document()
    doc.add_heading(title or "Use-case", level=1)
    for block in content.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        first_line = block.splitlines()[0]
        if ":" in first_line and len(first_line) < 80:
            heading = first_line.rstrip(":")
            rest = "\n".join(block.splitlines()[1:]).strip()
            doc.add_heading(heading, level=2)
            if rest:
                for line in rest.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line)
        else:
            for line in block.split("\n"):
                if line.strip():
                    doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ----------------------------
# Convenience function: combine alles en render string
# ----------------------------

def generate_rendered_template(short_input: str, template_name: str, templates_dir: str, use_groq: bool = True) -> Tuple[str, Dict[str, str]]:
    """
    Genereer mapping en render template. Retourneer (rendered_text, mapping).
    """
    templates = load_templates(templates_dir)
    tpl = templates.get(template_name)
    if tpl is None:
        raise FileNotFoundError(f"Template '{template_name}' niet gevonden in {templates_dir}.")

    mapping = generate_for_template(short_input, template_name, templates_dir, use_groq=use_groq)
    rendered = render_template(tpl, mapping)
    return rendered, mapping


# ----------------------------
# CLI / test snippet
# ----------------------------

if __name__ == "__main__":
    # eenvoudige lokale test
    here = os.path.dirname(__file__)
    templates_dir = os.path.join(here, "templates")
    # probeer 'story_template' zoals in jouw screenshot
    candidate_names = ["story_template", "story", "feature_template", "feature", "epic_template", "epic"]
    templates = load_templates(templates_dir)
    chosen = None
    for c in candidate_names:
        if c in templates:
            chosen = c
            break
    if not chosen:
        print("Geen template gevonden. Plaats story_template/feature_template/epic_template in templates/")
        print("Beschikbare:", list(templates.keys()))
        raise SystemExit(1)

    example_short = "Als gebruiker wil ik inloggen met eenmalige verificatie zodat ik veilig bij mijn documenten kan."
    rendered, mapping = generate_rendered_template(example_short, chosen, templates_dir, use_groq=False)
    print("---- MAPPING ----")
    print(json.dumps(mapping, indent=2, ensure_ascii=False))
    print("\n---- RENDERED ----\n")
    print(rendered)
