# core/docx_utils.py
import io
import tempfile
from typing import List, Dict
import docx
from docx.enum.text import WD_BREAK

def read_docx(path: str) -> str:
    doc = docx.Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def apply_replacements(doc_path: str, replacements: List[Dict], include_changes_overview: bool = True) -> bytes:
    doc = docx.Document(doc_path)
    def repl(runs):
        if not runs: return
        txt = "".join(r.text for r in runs)
        for rp in replacements:
            txt = txt.replace(rp["find"], rp["replace"])
        runs[0].text = txt
        for r in runs[1:]:
            r.text = ""
    for p in doc.paragraphs:
        repl(p.runs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    repl(p.runs)
    if include_changes_overview and replacements:
        try:
            p_br = doc.add_paragraph()
            run = p_br.add_run()
            run.add_break(WD_BREAK.PAGE)
        except Exception:
            pass
        try:
            doc.add_paragraph("Aangepaste onderdelen", style="Heading 1")
        except Exception:
            doc.add_paragraph("Aangepaste onderdelen")
        for rp in replacements:
            para = doc.add_paragraph(f"• {rp['find']} → {rp['replace']}")
            try:
                para.style = "List Bullet"
            except Exception:
                pass
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
